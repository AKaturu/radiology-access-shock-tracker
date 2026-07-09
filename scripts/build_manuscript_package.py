from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as ReportLabImage,
)
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DRAFT = ROOT / "docs" / "MANUSCRIPT_DRAFT.md"
DEFAULT_OUTPUT_DIR = ROOT / "dist" / "manuscript"
DOCX_NAME = "radiology-access-shock-tracker-manuscript.docx"
PDF_NAME = "radiology-access-shock-tracker-manuscript.pdf"

FIGURES = [
    (
        ROOT / "docs" / "assets" / "github" / "dashboard-overview.png",
        "Figure 1. Dashboard overview showing the publication-boundary-aware surveillance "
        "interface.",
    ),
    (
        ROOT / "docs" / "assets" / "github" / "readiness-audit.png",
        "Figure 2. Readiness audit view used to keep publication blockers visible.",
    ),
    (
        ROOT / "docs" / "assets" / "github" / "interventions.png",
        "Figure 3. Candidate intervention ranking output for access-recovery planning review.",
    ),
]

CITATION_PLACEHOLDERS = [
    (
        "[CITATION: FDA MQSA public facility file]",
        "Use for FDA MQSA source-file description and public-data limitations.",
    ),
    (
        "[CITATION: US Census ACS and Gazetteer]",
        "Use for county and tract population/context methods.",
    ),
    (
        "[CITATION: CDC PLACES and CDC/ATSDR SVI]",
        "Use for contextual vulnerability and mammography screening context.",
    ),
    (
        "[CITATION: HRSA service delivery sites]",
        "Use for candidate-site source assumptions and limitations.",
    ),
    (
        "[CITATION: OSRM/OpenStreetMap routing]",
        "Use for self-hosted route-time matrix methods and routing limitations.",
    ),
    (
        "[CITATION: Radiology Access Shock Tracker v0.2.0 release]",
        "Use for software release, reproducibility, checksum, and SBOM references.",
    ),
]

BOUNDARY_NOTE = (
    "Submission boundary: the current evidence supports software/methods claims, a reviewed NC "
    "row-level validation package, and 51-jurisdiction readiness claims. Replace citation "
    "placeholders before journal submission. Do not claim all-state row-level findings, clinical "
    "validation, confirmed closures, longitudinal access deterioration, or causal utilization "
    "effects until the required evidence exists."
)


@dataclass(frozen=True)
class MarkdownBlock:
    kind: str
    level: int
    text: str


def parse_markdown(path: Path) -> list[MarkdownBlock]:
    blocks: list[MarkdownBlock] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_lines:
            blocks.append(MarkdownBlock("paragraph", 0, " ".join(paragraph_lines)))
            paragraph_lines.clear()

    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            flush_paragraph()
            continue
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading_match:
            flush_paragraph()
            blocks.append(
                MarkdownBlock(
                    "heading",
                    len(heading_match.group(1)),
                    heading_match.group(2).strip(),
                )
            )
            continue
        if line.startswith("- "):
            flush_paragraph()
            blocks.append(MarkdownBlock("bullet", 0, line[2:].strip()))
            continue
        paragraph_lines.append(line)
    flush_paragraph()
    return blocks


def set_run_font(
    paragraph,
    *,
    size_pt: float,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
) -> None:
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color


def configure_docx_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_header_footer(document: Document) -> None:
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Radiology Access Shock Tracker - Manuscript Draft"
    set_run_font(header, size_pt=9, color=RGBColor(85, 85, 85))

    footer = section.footer.paragraphs[0]
    footer.text = "Submission draft - replace citation placeholders before submission"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer, size_pt=9, color=RGBColor(85, 85, 85))


def add_docx_title(document: Document) -> None:
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run(
        "Radiology Access Shock Tracker: A Reproducible Review-Gated Workflow for "
        "Mammography Access Surveillance"
    )
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(18)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run(
        "Submission-ready working draft with bounded claims, citation placeholders, and figures"
    )
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(11)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor(85, 85, 85)

    metadata = document.add_table(rows=3, cols=2)
    metadata.autofit = False
    metadata.columns[0].width = Inches(1.45)
    metadata.columns[1].width = Inches(5.05)
    rows = [
        ("Evidence scope", "Reviewed NC row-level validation; 51-jurisdiction readiness package"),
        ("Draft source", "docs/MANUSCRIPT_DRAFT.md"),
        ("Citation status", "Placeholders included; final journal references still required"),
    ]
    for row, (label, value) in zip(metadata.rows, rows, strict=True):
        row.cells[0].text = label
        row.cells[1].text = value
        for paragraph in row.cells[0].paragraphs:
            set_run_font(paragraph, size_pt=9.5, bold=True)
        for paragraph in row.cells[1].paragraphs:
            set_run_font(paragraph, size_pt=9.5)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(10)
    note_run = note.add_run(BOUNDARY_NOTE)
    note_run.bold = True
    note_run.font.name = "Calibri"
    note_run.font.size = Pt(10.5)
    note_run.font.color.rgb = RGBColor(31, 58, 95)


def add_docx_markdown(document: Document, blocks: list[MarkdownBlock]) -> None:
    skip_heading = "Manuscript Working Draft"
    for block in blocks:
        if block.kind == "heading":
            if block.text == skip_heading:
                continue
            style_name = "Heading 1" if block.level <= 2 else "Heading 2"
            if block.level >= 4:
                style_name = "Heading 3"
            document.add_paragraph(block.text, style=style_name)
        elif block.kind == "bullet":
            paragraph = document.add_paragraph(block.text, style="List Bullet")
            set_run_font(paragraph, size_pt=11)
        else:
            paragraph = document.add_paragraph(block.text)
            set_run_font(paragraph, size_pt=11)


def add_docx_figures(document: Document) -> None:
    document.add_paragraph("Figures", style="Heading 1")
    for figure_path, caption in FIGURES:
        if not figure_path.exists():
            continue
        document.add_picture(str(figure_path), width=Inches(6.0))
        caption_para = document.add_paragraph(caption)
        caption_para.paragraph_format.space_after = Pt(10)
        set_run_font(caption_para, size_pt=9.5, italic=True, color=RGBColor(85, 85, 85))


def add_docx_citations(document: Document) -> None:
    document.add_paragraph("Citation Placeholders", style="Heading 1")
    document.add_paragraph(
        "Replace these placeholders with journal-formatted references before submission."
    )
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.25)
    table.columns[1].width = Inches(4.25)
    table.rows[0].cells[0].text = "Placeholder"
    table.rows[0].cells[1].text = "Use"
    for placeholder, use in CITATION_PLACEHOLDERS:
        cells = table.add_row().cells
        cells[0].text = placeholder
        cells[1].text = use
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_run_font(paragraph, size_pt=9.5, bold=row_index == 0)


def build_docx(blocks: list[MarkdownBlock], output_path: Path) -> None:
    document = Document()
    configure_docx_styles(document)
    add_header_footer(document)
    add_docx_title(document)
    add_docx_markdown(document, blocks)
    add_docx_figures(document)
    add_docx_citations(document)
    document.sections[-1].start_type = WD_SECTION_START.NEW_PAGE
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def pdf_styles() -> dict[str, ParagraphStyle]:
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ManuscriptTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=17,
            leading=21,
            textColor=colors.HexColor("#0B2545"),
            alignment=TA_CENTER,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BoundaryNote",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=9.5,
            leading=12,
            textColor=colors.HexColor("#1F3A5F"),
            backColor=colors.HexColor("#F4F6F9"),
            borderColor=colors.HexColor("#D7DBE2"),
            borderWidth=0.5,
            borderPadding=7,
            spaceBefore=6,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CitationCell",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=7.8,
            leading=9.2,
            spaceAfter=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CitationHeader",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=8.2,
            leading=9.5,
            textColor=colors.HexColor("#0B2545"),
            spaceAfter=0,
        )
    )
    styles["BodyText"].fontName = "Helvetica"
    styles["BodyText"].fontSize = 10.5
    styles["BodyText"].leading = 13
    styles["BodyText"].spaceAfter = 6
    styles["Heading1"].textColor = colors.HexColor("#2E74B5")
    styles["Heading2"].textColor = colors.HexColor("#2E74B5")
    styles["Heading3"].textColor = colors.HexColor("#1F4D78")
    return styles


def add_pdf_image(
    story: list[object],
    image_path: Path,
    caption: str,
    styles: dict[str, ParagraphStyle],
) -> None:
    if not image_path.exists():
        return
    with PILImage.open(image_path) as image:
        width, height = image.size
    max_width = 6.2 * inch
    max_height = 3.25 * inch
    ratio = min(max_width / width, max_height / height, 1.0)
    story.append(
        KeepTogether(
            [
                ReportLabImage(str(image_path), width=width * ratio, height=height * ratio),
                Paragraph(caption, styles["Italic"]),
                Spacer(1, 0.14 * inch),
            ]
        )
    )


def build_pdf(blocks: list[MarkdownBlock], output_path: Path) -> None:
    styles = pdf_styles()
    story: list[object] = [
        Paragraph(
            "Radiology Access Shock Tracker: A Reproducible Review-Gated Workflow for "
            "Mammography Access Surveillance",
            styles["ManuscriptTitle"],
        ),
        Paragraph(
            "Submission-ready working draft with bounded claims, citation placeholders, "
            "and figures",
            styles["Italic"],
        ),
        Spacer(1, 0.12 * inch),
        Paragraph(BOUNDARY_NOTE, styles["BoundaryNote"]),
    ]

    for block in blocks:
        if block.kind == "heading":
            if block.text == "Manuscript Working Draft":
                continue
            style_name = "Heading1" if block.level <= 2 else "Heading2"
            if block.level >= 4:
                style_name = "Heading3"
            story.append(Paragraph(block.text, styles[style_name]))
        elif block.kind == "bullet":
            story.append(Paragraph(block.text, styles["BodyText"], bulletText="-"))
        else:
            story.append(Paragraph(block.text, styles["BodyText"]))

    story.append(PageBreak())
    story.append(Paragraph("Figures", styles["Heading1"]))
    for image_path, caption in FIGURES:
        add_pdf_image(story, image_path, caption, styles)

    story.append(Paragraph("Citation Placeholders", styles["Heading1"]))
    citation_rows = [
        [
            Paragraph("Placeholder", styles["CitationHeader"]),
            Paragraph("Use", styles["CitationHeader"]),
        ],
        *[
            [
                Paragraph(placeholder, styles["CitationCell"]),
                Paragraph(use, styles["CitationCell"]),
            ]
            for placeholder, use in CITATION_PLACEHOLDERS
        ],
    ]
    citation_table = Table(citation_rows, colWidths=[2.7 * inch, 3.6 * inch])
    citation_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2545")),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#D7DBE2")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(citation_table)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    pdf = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=72,
        rightMargin=72,
        topMargin=72,
        bottomMargin=72,
        title="Radiology Access Shock Tracker Manuscript",
    )
    pdf.build(story)


def build_package(draft: Path, output_dir: Path) -> tuple[Path, Path]:
    blocks = parse_markdown(draft)
    docx_path = output_dir / DOCX_NAME
    pdf_path = output_dir / PDF_NAME
    build_docx(blocks, docx_path)
    build_pdf(blocks, pdf_path)
    return docx_path, pdf_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Build the manuscript DOCX/PDF package.")
    parser.add_argument("--draft", type=Path, default=DEFAULT_DRAFT)
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR)
    args = parser.parse_args()

    docx_path, pdf_path = build_package(args.draft, args.output_dir)
    print(f"Wrote {docx_path}")
    print(f"Wrote {pdf_path}")


if __name__ == "__main__":
    main()
